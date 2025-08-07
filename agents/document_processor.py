import logging
from typing import Dict, Any, List, Optional, Tuple
import os
from pathlib import Path
import asyncio
from datetime import datetime
import base64
import io
import json
import aiohttp

# Document processing libraries
import pymupdf4llm  # PDF to markdown conversion
import fitz  # PyMuPDF for additional PDF operations
import openpyxl  # Excel
from PIL import Image  # Images
import ezdxf  # AutoCAD DWG/DXF
from docx import Document as WordDoc  # Word documents

from agents.base_agent import BaseAgent
from storage import mongodb_connector, faiss_connector
from models import Document, DocumentType, DocumentStatus, DocumentMetadata, ProcessedData
from config.settings import config
from services.file_storage_service import file_storage, FileType, FileAction

logger = logging.getLogger(__name__)

class DocumentProcessorAgent(BaseAgent):
    """
    The Eyes of the system - understands all types of documents and images
    Uses Gemini 2.5 Flash for native OCR and document understanding
    Uses pymupdf4llm for efficient PDF to markdown conversion
    """
    
    def __init__(self):
        super().__init__(
            name="document_processor",
            description="Processes all document types with Gemini 2.5 Flash OCR and pymupdf4llm"
        )
        self.supported_formats = {
            ".pdf": DocumentType.PDF,
            ".xlsx": DocumentType.EXCEL,
            ".xls": DocumentType.EXCEL,
            ".docx": DocumentType.WORD,
            ".doc": DocumentType.WORD,
            ".dwg": DocumentType.AUTOCAD,
            ".dxf": DocumentType.AUTOCAD,
            ".png": DocumentType.IMAGE,
            ".jpg": DocumentType.IMAGE,
            ".jpeg": DocumentType.IMAGE,
            ".bmp": DocumentType.IMAGE,
            ".tiff": DocumentType.IMAGE
        }
        self.gemini_api_key = config.OPENROUTER_API_KEY
        self.gemini_model = config.AI_MODEL
    
    async def validate_request(self, request: Dict[str, Any]) -> bool:
        """Validate document processing request"""
        return "file_path" in request or "document_id" in request
    
    async def process(self, request: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process document based on type
        """
        try:
            context = request.get("context", {})
            parameters = request.get("parameters", {})
            
            # Get document info
            file_path = request.get("file_path")
            document_id = request.get("document_id")
            
            if document_id:
                # Load existing document
                doc = await mongodb_connector.get_document(document_id)
                if not doc:
                    raise ValueError(f"Document {document_id} not found")
                file_path = doc["file_path"]
            
            # Detect document type
            doc_type = self._detect_document_type(file_path)
            
            # Create document record
            document = await self._create_document_record(
                file_path,
                doc_type,
                context
            )
            
            # Process based on type
            if doc_type == DocumentType.PDF:
                processed_data = await self._process_pdf_with_pymupdf4llm(file_path, parameters)
            elif doc_type == DocumentType.EXCEL:
                processed_data = await self._process_excel(file_path, parameters)
            elif doc_type == DocumentType.WORD:
                processed_data = await self._process_word(file_path, parameters)
            elif doc_type == DocumentType.AUTOCAD:
                processed_data = await self._process_autocad(file_path, parameters)
            elif doc_type == DocumentType.IMAGE:
                processed_data = await self._process_image_with_gemini(file_path, parameters)
            else:
                processed_data = await self._process_text(file_path, parameters)
            
            # Update document with processed data
            document.processed_data = processed_data
            document.status = DocumentStatus.COMPLETED
            document.processed_at = datetime.utcnow()
            
            # Save to database
            doc_id = await mongodb_connector.save_document(document.dict())
            document.id = doc_id
            
            # Generate embeddings if requested
            if parameters.get("generate_embeddings", True):
                await self._generate_document_embeddings(document)
            
            # Format response
            response_data = self._format_response(document, processed_data)
            
            return {
                "data": response_data,
                "metadata": {
                    "document_id": doc_id,
                    "processing_time": (datetime.utcnow() - document.created_at).total_seconds(),
                    "document_type": doc_type.value,
                    "processor": "gemini_2.5_flash" if doc_type in [DocumentType.PDF, DocumentType.IMAGE] else "native"
                }
            }
            
        except Exception as e:
            logger.error(f"Document processing error: {e}", exc_info=True)
            raise
    
    def _detect_document_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension"""
        ext = Path(file_path).suffix.lower()
        return self.supported_formats.get(ext, DocumentType.UNKNOWN)
    
    async def _create_document_record(
        self,
        file_path: str,
        doc_type: DocumentType,
        context: Dict[str, Any]
    ) -> Document:
        """Create document record in database with new file storage"""
        file_stats = os.stat(file_path)
        user_id = context.get("user_id", 0)
        
        # Read file content
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        # Determine file type for storage
        storage_file_type = FileType.DOCUMENT
        if doc_type == DocumentType.IMAGE:
            storage_file_type = FileType.IMAGE
        elif doc_type in [DocumentType.PDF, DocumentType.EXCEL, DocumentType.WORD]:
            storage_file_type = FileType.DOCUMENT
        elif doc_type == DocumentType.AUTOCAD:
            storage_file_type = FileType.DATA
        
        # Save to new file storage structure
        storage_info = file_storage.save_upload(
            user_id=user_id,
            file_content=file_content,
            filename=Path(file_path).name,
            file_type=storage_file_type,
            metadata={
                "document_type": doc_type.value,
                "project_id": context.get("project_id"),
                "client_id": context.get("client_id")
            }
        )
        
        metadata = DocumentMetadata(
            filename=Path(file_path).name,
            size_bytes=file_stats.st_size,
            mime_type=self._get_mime_type(doc_type)
        )
        
        document = Document(
            id="",  # Will be set by MongoDB
            user_id=user_id,
            project_id=context.get("project_id"),
            file_path=storage_info["storage_path"],  # Use new storage path
            original_name=Path(file_path).name,
            document_type=doc_type,
            status=DocumentStatus.PROCESSING,
            metadata=metadata
        )
        
        return document
    
    async def _process_pdf_with_pymupdf4llm(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> ProcessedData:
        """
        Process PDF document using pymupdf4llm for LLM-ready markdown
        and Gemini 2.5 Flash for OCR if needed
        """
        try:
            # Use pymupdf4llm to convert PDF to markdown
            markdown_text = pymupdf4llm.to_markdown(
                file_path,
                page_chunks=True,  # Separate pages
                write_images=True,  # Extract images
                image_path="temp_images",  # Temporary image storage
                dpi=150  # Good quality for OCR
            )
            
            # Process with chunks for large PDFs
            chunks = []
            tables = []
            images = []
            
            # If the PDF has images or scanned content, use Gemini for OCR
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            # Check if OCR is needed by looking for text
            needs_ocr = False
            for page in doc:
                if not page.get_text().strip():
                    needs_ocr = True
                    break
            
            if needs_ocr and parameters.get("ocr_if_needed", True):
                # Process in batches with Gemini (1,500 pages at once capability)
                batch_size = min(parameters.get("batch_size", 100), 1500)
                
                for start_idx in range(0, total_pages, batch_size):
                    end_idx = min(start_idx + batch_size, total_pages)
                    
                    # Extract pages as images for OCR
                    page_images = []
                    for page_num in range(start_idx, end_idx):
                        page = doc[page_num]
                        pix = page.get_pixmap(dpi=150)
                        img_data = pix.tobytes("png")
                        img_base64 = base64.b64encode(img_data).decode()
                        page_images.append(img_base64)
                    
                    # Send batch to Gemini for OCR
                    ocr_result = await self._batch_ocr_with_gemini(
                        page_images,
                        f"Pages {start_idx + 1} to {end_idx} of PDF"
                    )
                    
                    chunks.append(ocr_result)
            else:
                # Use the markdown text from pymupdf4llm
                chunks = [markdown_text]
            
            # Extract tables if present
            if parameters.get("extract_tables", True):
                # pymupdf4llm preserves table structure in markdown
                tables = self._extract_tables_from_markdown(markdown_text)
            
            # Extract images
            image_count = 0
            for page_num, page in enumerate(doc):
                imgs = self._extract_pdf_images(page, page_num)
                images.extend(imgs)
                image_count += len(imgs)
            
            doc.close()
            
            # Clean up temporary images
            if os.path.exists("temp_images"):
                import shutil
                shutil.rmtree("temp_images")
            
            return ProcessedData(
                text_content="\n\n".join(chunks) if chunks else markdown_text,
                tables=tables,
                images=images[:50],  # Limit to 50 images to avoid memory issues
                structured_data={
                    "page_count": total_pages,
                    "has_ocr": needs_ocr,
                    "has_tables": len(tables) > 0,
                    "has_images": image_count > 0,
                    "image_count": image_count,
                    "processor": "pymupdf4llm + gemini" if needs_ocr else "pymupdf4llm"
                }
            )
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise
    
    async def _batch_ocr_with_gemini(
        self,
        images: List[str],
        context: str
    ) -> str:
        """
        Perform OCR on batch of images using Gemini 2.5 Flash
        Can handle up to 1,500 pages in one request
        """
        try:
            # Prepare the multimodal request for Gemini
            messages = [
                {
                    "role": "system",
                    "content": "You are an expert OCR system for construction documents. Extract all text, preserving structure, tables, and formatting. Return the content in markdown format."
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"Please perform OCR on these {len(images)} pages from {context}. Extract all text, tables, and maintain document structure. Format the output as markdown."
                        }
                    ]
                }
            ]
            
            # Add images to the request
            for i, img_base64 in enumerate(images):
                messages[1]["content"].append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{img_base64}"
                    }
                })
            
            # Call Gemini via OpenRouter
            headers = {
                "Authorization": f"Bearer {self.gemini_api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://spencerai.com",
                "X-Title": "Spencer AI Document Processor"
            }
            
            payload = {
                "model": self.gemini_model,
                "messages": messages,
                "temperature": 0.1,  # Low temperature for accuracy
                "max_tokens": 100000  # Gemini 2.5 Flash supports large outputs
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{config.OPENROUTER_BASE_URL}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=300)  # 5 minute timeout
                ) as response:
                    if response.status == 200:
                        result = await response.json()
                        return result["choices"][0]["message"]["content"]
                    else:
                        error = await response.text()
                        logger.error(f"Gemini OCR error: {error}")
                        return f"OCR failed: {error}"
                        
        except Exception as e:
            logger.error(f"Gemini batch OCR error: {e}")
            return f"OCR error: {str(e)}"
    
    async def _process_excel(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> ProcessedData:
        """Process Excel spreadsheet"""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            tables = []
            text_content = []
            structured_data = {
                "sheets": [],
                "total_rows": 0,
                "total_cells": 0
            }
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # Extract data as table
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append(list(row))
                
                if sheet_data:
                    tables.append({
                        "name": sheet_name,
                        "headers": sheet_data[0] if sheet_data else [],
                        "rows": sheet_data[1:] if len(sheet_data) > 1 else [],
                        "row_count": len(sheet_data)
                    })
                    
                    structured_data["sheets"].append({
                        "name": sheet_name,
                        "rows": len(sheet_data),
                        "columns": len(sheet_data[0]) if sheet_data else 0
                    })
                    structured_data["total_rows"] += len(sheet_data)
                
                # Extract formulas and comments if present
                text_content.append(f"Sheet: {sheet_name}")
                
            wb.close()
            
            # If requested, use Gemini to analyze the spreadsheet data
            if parameters.get("analyze_with_ai", False) and tables:
                analysis = await self._analyze_spreadsheet_with_gemini(tables)
                text_content.append(f"\nAI Analysis:\n{analysis}")
            
            return ProcessedData(
                text_content="\n".join(text_content),
                tables=tables,
                structured_data=structured_data
            )
            
        except Exception as e:
            logger.error(f"Excel processing error: {e}")
            raise
    
    async def _process_autocad(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> ProcessedData:
        """Process AutoCAD DWG/DXF files"""
        try:
            # Read DXF file (DWG would need conversion first)
            doc = ezdxf.readfile(file_path)
            
            text_content = []
            structured_data = {
                "layers": [],
                "blocks": [],
                "dimensions": [],
                "entities": {}
            }
            
            # Extract layers
            for layer in doc.layers:
                layer_info = {
                    "name": layer.dxf.name,
                    "color": layer.dxf.color,
                    "linetype": layer.dxf.linetype,
                    "on": not layer.is_off()
                }
                structured_data["layers"].append(layer_info)
                text_content.append(f"Layer: {layer.dxf.name}")
            
            # Extract blocks (reusable components)
            for block in doc.blocks:
                if block.name.startswith("*"):  # Skip anonymous blocks
                    continue
                structured_data["blocks"].append({
                    "name": block.name,
                    "entity_count": len(block)
                })
            
            # Count entities by type
            msp = doc.modelspace()
            entity_counts = {}
            
            for entity in msp:
                entity_type = entity.dxftype()
                entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
                
                # Extract dimensions
                if entity_type == "DIMENSION":
                    structured_data["dimensions"].append({
                        "type": entity.dxf.dimtype,
                        "text": entity.dxf.text
                    })
                    text_content.append(f"Dimension: {entity.dxf.text}")
                
                # Extract text entities
                elif entity_type in ["TEXT", "MTEXT"]:
                    text_content.append(f"Text: {entity.dxf.text}")
            
            structured_data["entities"] = entity_counts
            
            return ProcessedData(
                text_content="\n".join(text_content),
                structured_data=structured_data
            )
            
        except Exception as e:
            logger.error(f"AutoCAD processing error: {e}")
            raise
    
    async def _process_image_with_gemini(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> ProcessedData:
        """Process image files with Gemini 2.5 Flash multimodal capabilities"""
        try:
            img = Image.open(file_path)
            
            # Get image metadata
            structured_data = {
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode,
                "dpi": img.info.get("dpi", (72, 72))
            }
            
            # Convert to base64 for Gemini
            buffered = io.BytesIO()
            img.save(buffered, format=img.format or "PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode()
            
            # Use Gemini for image analysis with construction context
            analysis_prompt = parameters.get("analysis_prompt", 
                "Analyze this construction/engineering image. Describe what you see, "
                "identify any text, measurements, equipment, materials, safety concerns, "
                "or technical details. If this is a blueprint or technical drawing, "
                "describe the components and any visible specifications."
            )
            
            # Call Gemini for analysis
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": analysis_prompt
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/{img.format or 'png'};base64,{img_base64}"
                            }
                        }
                    ]
                }
            ]
            
            headers = {
                "Authorization": f"Bearer {self.gemini_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": self.gemini_model,
                "messages": messages,
                "temperature": 0.3,
                "max_tokens": 2000
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{config.OPENROUTER_BASE_URL}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status == 200:
                        result = await response.json()
                        image_description = result["choices"][0]["message"]["content"]
                    else:
                        image_description = "Failed to analyze image with Gemini"
            
            return ProcessedData(
                text_content=image_description,
                images=[img_base64],
                structured_data=structured_data
            )
            
        except Exception as e:
            logger.error(f"Image processing error: {e}")
            raise
    
    async def _process_word(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> ProcessedData:
        """Process Word documents"""
        try:
            doc = WordDoc(file_path)
            
            text_content = []
            tables = []
            images = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "name": f"Table {i + 1}",
                        "headers": table_data[0] if table_data else [],
                        "rows": table_data[1:] if len(table_data) > 1 else []
                    })
            
            return ProcessedData(
                text_content="\n\n".join(text_content),
                tables=tables,
                structured_data={
                    "paragraph_count": len(text_content),
                    "table_count": len(tables),
                    "sections": len(doc.sections)
                }
            )
            
        except Exception as e:
            logger.error(f"Word processing error: {e}")
            raise
    
    async def _process_text(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> ProcessedData:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return ProcessedData(
                text_content=content,
                structured_data={
                    "line_count": len(content.splitlines()),
                    "character_count": len(content)
                }
            )
            
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            raise
    
    def _extract_tables_from_markdown(self, markdown_text: str) -> List[Dict[str, Any]]:
        """Extract tables from markdown text generated by pymupdf4llm"""
        tables = []
        lines = markdown_text.split('\n')
        current_table = []
        in_table = False
        
        for line in lines:
            # Detect table rows (contain |)
            if '|' in line and not line.strip().startswith('```'):
                in_table = True
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                current_table.append(cells)
            elif in_table and '|' not in line:
                # End of table
                if current_table:
                    # First row is headers, rest are data
                    tables.append({
                        "name": f"Table {len(tables) + 1}",
                        "headers": current_table[0] if current_table else [],
                        "rows": current_table[2:] if len(current_table) > 2 else []  # Skip separator row
                    })
                current_table = []
                in_table = False
        
        # Handle last table if exists
        if current_table:
            tables.append({
                "name": f"Table {len(tables) + 1}",
                "headers": current_table[0] if current_table else [],
                "rows": current_table[2:] if len(current_table) > 2 else []
            })
        
        return tables
    
    def _extract_pdf_images(self, page, page_num: int) -> List[str]:
        """Extract images from PDF page"""
        images = []
        
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            try:
                # Get image data
                xref = img[0]
                pix = fitz.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    img_base64 = base64.b64encode(img_data).decode()
                    images.append(img_base64)
                
                pix = None
                
            except Exception as e:
                logger.error(f"Error extracting image: {e}")
        
        return images
    
    async def _analyze_spreadsheet_with_gemini(self, tables: List[Dict]) -> str:
        """Use Gemini to analyze spreadsheet data"""
        try:
            # Prepare data summary for Gemini
            summary = "Spreadsheet Analysis Request:\n\n"
            
            for table in tables[:5]:  # Limit to first 5 sheets
                summary += f"Sheet: {table['name']}\n"
                summary += f"Headers: {', '.join(str(h) for h in table['headers'][:10])}\n"
                summary += f"Row count: {table['row_count']}\n"
                summary += f"Sample data (first 5 rows):\n"
                
                for row in table['rows'][:5]:
                    summary += f"  {row[:10]}\n"  # First 10 columns
                summary += "\n"
            
            messages = [
                {
                    "role": "user",
                    "content": f"{summary}\n\nPlease analyze this construction/engineering spreadsheet data. Identify key metrics, patterns, and any insights relevant to construction project management."
                }
            ]
            
            headers = {
                "Authorization": f"Bearer {self.gemini_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": self.gemini_model,
                "messages": messages,
                "temperature": 0.3,
                "max_tokens": 1000
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{config.OPENROUTER_BASE_URL}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=aiohttp.ClientTimeout(total=30)
                ) as response:
                    if response.status == 200:
                        result = await response.json()
                        return result["choices"][0]["message"]["content"]
                    else:
                        return "Failed to analyze spreadsheet with AI"
                        
        except Exception as e:
            logger.error(f"Spreadsheet analysis error: {e}")
            return "Error analyzing spreadsheet"
    
    async def _generate_document_embeddings(self, document: Document):
        """Generate embeddings for document content"""
        try:
            # Get semantic search agent
            from agents.semantic_search import semantic_search
            
            # Prepare content for embedding
            content = document.processed_data.text_content or ""
            
            # Add structured data as text
            if document.processed_data.structured_data:
                content += f"\n\nMetadata: {document.processed_data.structured_data}"
            
            # Create chunks if content is too long
            chunks = self._chunk_text(content, max_length=1000)
            
            # Index each chunk
            for i, chunk in enumerate(chunks):
                await semantic_search._index_content({
                    "content": {
                        "id": f"{document.id}_chunk_{i}",
                        "text": chunk,
                        "metadata": {
                            "document_id": document.id,
                            "document_type": document.document_type.value,
                            "filename": document.original_name,
                            "chunk_index": i,
                            "total_chunks": len(chunks)
                        }
                    },
                    "context": {"user_id": document.user_id}
                })
            
            document.processed_data.embeddings_generated = True
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
    
    def _chunk_text(self, text: str, max_length: int = 1000) -> List[str]:
        """Split text into chunks for embedding"""
        chunks = []
        current_chunk = ""
        
        for sentence in text.split(". "):
            if len(current_chunk) + len(sentence) < max_length:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _format_response(
        self,
        document: Document,
        processed_data: ProcessedData
    ) -> Dict[str, Any]:
        """Format processing response"""
        response = {
            "document_id": document.id,
            "filename": document.original_name,
            "document_type": document.document_type.value,
            "status": document.status.value,
            "summary": self._generate_summary(document, processed_data),
            "metadata": document.metadata.dict()
        }
        
        # Add type-specific data
        if document.document_type == DocumentType.EXCEL:
            response["sheets"] = processed_data.structured_data.get("sheets", [])
            response["total_rows"] = processed_data.structured_data.get("total_rows", 0)
            
        elif document.document_type == DocumentType.AUTOCAD:
            response["layers"] = processed_data.structured_data.get("layers", [])
            response["entity_count"] = sum(
                processed_data.structured_data.get("entities", {}).values()
            )
            
        elif document.document_type == DocumentType.PDF:
            response["page_count"] = processed_data.structured_data.get("page_count", 0)
            response["has_images"] = processed_data.structured_data.get("has_images", False)
            response["has_ocr"] = processed_data.structured_data.get("has_ocr", False)
            response["processor"] = processed_data.structured_data.get("processor", "unknown")
        
        # Add tables if present
        if processed_data.tables:
            response["tables"] = [
                {
                    "name": t.get("name", ""),
                    "row_count": len(t.get("rows", [])),
                    "column_count": len(t.get("headers", []))
                }
                for t in processed_data.tables
            ]
        
        return response
    
    def _generate_summary(
        self,
        document: Document,
        processed_data: ProcessedData
    ) -> str:
        """Generate document summary"""
        doc_type = document.document_type.value
        
        if doc_type == "pdf":
            pages = processed_data.structured_data.get("page_count", 0)
            processor = processed_data.structured_data.get("processor", "")
            ocr_used = processed_data.structured_data.get("has_ocr", False)
            return f"PDF document with {pages} pages processed using {processor}" + (" with OCR" if ocr_used else "")
            
        elif doc_type == "excel":
            sheets = len(processed_data.structured_data.get("sheets", []))
            rows = processed_data.structured_data.get("total_rows", 0)
            return f"Excel spreadsheet with {sheets} sheets and {rows} total rows"
            
        elif doc_type == "autocad":
            layers = len(processed_data.structured_data.get("layers", []))
            return f"AutoCAD drawing with {layers} layers analyzed"
            
        elif doc_type == "image":
            dims = processed_data.structured_data
            return f"Image ({dims.get('width')}x{dims.get('height')}) analyzed with Gemini 2.5 Flash"
            
        else:
            return f"{doc_type.title()} document processed successfully"
    
    def _get_mime_type(self, doc_type: DocumentType) -> str:
        """Get MIME type for document type"""
        mime_types = {
            DocumentType.PDF: "application/pdf",
            DocumentType.EXCEL: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            DocumentType.WORD: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            DocumentType.AUTOCAD: "application/acad",
            DocumentType.IMAGE: "image/png",
            DocumentType.TEXT: "text/plain"
        }
        return mime_types.get(doc_type, "application/octet-stream")

# Create singleton instance
document_processor = DocumentProcessorAgent()